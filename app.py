import streamlit as st
import os
import tempfile
from main import load_github_content, load_document, generate_model_card, customer_template, internal_template
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import re

# Streamlit app
st.set_page_config(page_title="AI Explainability Card Generator", layout="wide")

# Custom CSS for theme-agnostic design with correct button colors
st.markdown("""
    <style>
    .stButton > button {
        border-radius: 4px;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }
    .stButton > button:first-child {
        background-color: #1E90FF;
        color: white;
    }
    .stDownloadButton > button {
        background-color: #808080;
        color: white;
    }
    .column-content {
        padding: 1rem;
        overflow-y: auto;
        max-height: 70vh;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("AI Explainability Card Generator")

# Input fields
model_name = st.text_input("Model Name", "")
github_repo = st.text_input("GitHub Repository URL (optional)", "")
uploaded_files = st.file_uploader("Upload Documents (optional)", accept_multiple_files=True)

# Save as docx function
def save_as_docx(content, filename):
    doc = Document()

    # Function to get or create style
    def get_or_create_style(styles, name, font_size, bold=False):
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(font_size)
        style.font.bold = bold
        return style

    # Add or get styles
    styles = doc.styles
    get_or_create_style(styles, 'Heading 1', 18, bold=True)
    get_or_create_style(styles, 'Heading 2', 16, bold=True)
    get_or_create_style(styles, 'Normal', 11)

    # Parse the markdown content and add to document
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_paragraph(line[2:], style='Heading 1')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 2')
        elif line.startswith('- '):
            # Handle bullet points with potential bold text
            paragraph = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', line[2:])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
        else:
            # Handle normal text with potential bold text
            paragraph = doc.add_paragraph(style='Normal')
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)

    doc.save(filename)
    return filename

# Generate button
if st.button("Generate Explainability Cards", key="generate_btn", help="Click to generate explainability cards"):
    if not model_name:
        st.error("Please provide a name for the explainability card.")
    elif not github_repo and not uploaded_files:
        st.error("Please provide either a GitHub repository or upload documents.")
    else:
        with st.spinner("Generating explainability cards..."):
            content = []

            # Load GitHub content
            if github_repo:
                try:
                    github_content = load_github_content(github_repo)
                    content.append(github_content)
                except Exception as e:
                    st.error(f"Error loading GitHub content: {str(e)}")

            # Load uploaded documents
            if uploaded_files:
                for uploaded_file in uploaded_files:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        tmp_file_path = tmp_file.name

                    try:
                        doc_content = load_document(tmp_file_path)
                        content.append(doc_content)
                    except ValueError as e:
                        st.error(f"Error loading {uploaded_file.name}: {str(e)}")
                    finally:
                        os.unlink(tmp_file_path)

            if content:
                all_content = "\n\n".join(content)

                st.session_state.customer_card = generate_model_card(all_content, customer_template)
                st.session_state.internal_card = generate_model_card(all_content, internal_template)
                st.session_state.model_name = model_name
            else:
                st.error("No valid content was processed. Please check your inputs and try again.")

# Display cards if they exist in session state
if 'customer_card' in st.session_state and 'internal_card' in st.session_state:
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Customer Explainability Card")
        st.markdown('<div class="column-content">', unsafe_allow_html=True)
        st.markdown(st.session_state.customer_card)
        st.markdown('</div>', unsafe_allow_html=True)

        # Create download button for customer card
        customer_docx = save_as_docx(st.session_state.customer_card, f"{st.session_state.model_name}_customer_explainability_card.docx")
        with open(customer_docx, "rb") as f:
            st.download_button(
                "Download Customer Explainability Card (DOCX)",
                f,
                file_name=f"{st.session_state.model_name}_customer_explainability_card.docx",
                key="download_customer",
                help="Click to download the customer explainability card as a DOCX file"
            )

    with col2:
        st.subheader("Internal Stakeholder Explainability Card")
        st.markdown('<div class="column-content">', unsafe_allow_html=True)
        st.markdown(st.session_state.internal_card)
        st.markdown('</div>', unsafe_allow_html=True)

        # Create download button for internal card
        internal_docx = save_as_docx(st.session_state.internal_card, f"{st.session_state.model_name}_internal_explainability_card.docx")
        with open(internal_docx, "rb") as f:
            st.download_button(
                "Download Internal Explainability Card (DOCX)",
                f,
                file_name=f"{st.session_state.model_name}_internal_explainability_card.docx",
                key="download_internal",
                help="Click to download the internal explainability card as a DOCX file"
            )

st.sidebar.header("About")
st.sidebar.info("This app generates AI explainability cards based on GitHub repositories and/or uploaded documents. It creates both customer-facing and internal stakeholder versions of the explainability card.")
